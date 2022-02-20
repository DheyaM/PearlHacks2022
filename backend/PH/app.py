import pymongo

# https://www.fullstackpython.com/flask-helpers-make-response-examples.html
from flask import Flask, render_template,  request, jsonify, make_response, session, flash, redirect
from flask_pymongo import PyMongo
import flask
import os

import docx
from docx import Document
import json
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from google.cloud import translate

app = Flask(__name__, static_url_path='')

UPLOAD_FOLDER = './upload'
ALLOWED_EXTENSIONS = {'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

map = {'aa': 'Afar', 'ab': 'Abkhazian', 'af': 'Afrikaans', 'ak': 'Akan', 'sq': 'Albanian', 'am': 'Amharic', 'ar': 'Arabic', 'an': 'Aragonese', 'hy': 'Armenian', 'as': 'Assamese', 'av': 'Avaric', 'ae': 'Avestan', 'ay': 'Aymara', 'az': 'Azerbaijani', 'ba': 'Bashkir', 'bm': 'Bambara', 'eu': 'Basque', 'be': 'Belarusian', 'bn': 'Bengali', 'bh': 'Bihari languages', 'bi': 'Bislama', 'bo': 'Tibetan', 'bs': 'Bosnian', 'br': 'Breton', 'bg': 'Bulgarian', 'my': 'Burmese', 'ca': 'Catalan; Valencian', 'cs': 'Czech', 'ch': 'Chamorro', 'ce': 'Chechen', 'zh': 'Chinese', 'cu': 'Church Slavic; Old Slavonic; Church Slavonic; Old Bulgarian; Old Church Slavonic', 'cv': 'Chuvash', 'kw': 'Cornish', 'co': 'Corsican', 'cr': 'Cree', 'cy': 'Welsh', 'da': 'Danish', 'de': 'German', 'dv': 'Divehi; Dhivehi; Maldivian', 'nl': 'Dutch; Flemish', 'dz': 'Dzongkha', 'el': 'Greek, Modern (1453-)', 'en': 'English', 'eo': 'Esperanto', 'et': 'Estonian', 'ee': 'Ewe', 'fo': 'Faroese', 'fa': 'Persian', 'fj': 'Fijian', 'fi': 'Finnish', 'fr': 'French', 'fy': 'Western Frisian', 'ff': 'Fulah', 'Ga': 'Georgian', 'gd': 'Gaelic; Scottish Gaelic', 'ga': 'Irish', 'gl': 'Galician', 'gv': 'Manx', 'gn': 'Guarani', 'gu': 'Gujarati', 'ht': 'Haitian; Haitian Creole', 'ha': 'Hausa', 'he': 'Hebrew', 'hz': 'Herero', 'hi': 'Hindi', 'ho': 'Hiri Motu', 'hr': 'Croatian', 'hu': 'Hungarian', 'ig': 'Igbo', 'is': 'Icelandic', 'io': 'Ido', 'ii': 'Sichuan Yi; Nuosu', 'iu': 'Inuktitut', 'ie': 'Interlingue; Occidental', 'ia': 'Interlingua (International Auxiliary Language Association)', 'id': 'Indonesian', 'ik': 'Inupiaq', 'it': 'Italian', 'jv': 'Javanese', 'ja': 'Japanese', 'kl': 'Kalaallisut; Greenlandic', 'kn': 'Kannada', 'ks': 'Kashmiri', 'ka': 'Georgian', 'kr': 'Kanuri', 'kk': 'Kazakh', 'km': 'Central Khmer', 'ki': 'Kikuyu; Gikuyu', 'rw': 'Kinyarwanda', 'ky': 'Kirghiz; Kyrgyz', 'kv': 'Komi',
       'kg': 'Kongo', 'ko': 'Korean', 'kj': 'Kuanyama; Kwanyama', 'ku': 'Kurdish', 'lo': 'Lao', 'la': 'Latin', 'lv': 'Latvian', 'li': 'Limburgan; Limburger; Limburgish', 'ln': 'Lingala', 'lt': 'Lithuanian', 'lb': 'Luxembourgish; Letzeburgesch', 'lu': 'Luba-Katanga', 'lg': 'Ganda', 'mk': 'Macedonian', 'mh': 'Marshallese', 'ml': 'Malayalam', 'mi': 'Maori', 'mr': 'Marathi', 'ms': 'Malay', 'Mi': 'Micmac', 'mg': 'Malagasy', 'mt': 'Maltese', 'mn': 'Mongolian', 'na': 'Nauru', 'nv': 'Navajo; Navaho', 'nr': 'Ndebele, South; South Ndebele', 'nd': 'Ndebele, North; North Ndebele', 'ng': 'Ndonga', 'ne': 'Nepali', 'nn': 'Norwegian Nynorsk; Nynorsk, Norwegian', 'nb': 'Bokmål, Norwegian; Norwegian Bokmål', 'no': 'Norwegian', 'oc': 'Occitan (post 1500)', 'oj': 'Ojibwa', 'or': 'Oriya', 'om': 'Oromo', 'os': 'Ossetian; Ossetic', 'pa': 'Panjabi; Punjabi', 'pi': 'Pali', 'pl': 'Polish', 'pt': 'Portuguese', 'ps': 'Pushto; Pashto', 'qu': 'Quechua', 'rm': 'Romansh', 'ro': 'Romanian; Moldavian; Moldovan', 'rn': 'Rundi', 'ru': 'Russian', 'sg': 'Sango', 'sa': 'Sanskrit', 'si': 'Sinhala; Sinhalese', 'sk': 'Slovak', 'sl': 'Slovenian', 'se': 'Northern Sami', 'sm': 'Samoan', 'sn': 'Shona', 'sd': 'Sindhi', 'so': 'Somali', 'st': 'Sotho, Southern', 'es': 'Spanish; Castilian', 'sc': 'Sardinian', 'sr': 'Serbian', 'ss': 'Swati', 'su': 'Sundanese', 'sw': 'Swahili', 'sv': 'Swedish', 'ty': 'Tahitian', 'ta': 'Tamil', 'tt': 'Tatar', 'te': 'Telugu', 'tg': 'Tajik', 'tl': 'Tagalog', 'th': 'Thai', 'ti': 'Tigrinya', 'to': 'Tonga (Tonga Islands)', 'tn': 'Tswana', 'ts': 'Tsonga', 'tk': 'Turkmen', 'tr': 'Turkish', 'tw': 'Twi', 'ug': 'Uighur; Uyghur', 'uk': 'Ukrainian', 'ur': 'Urdu', 'uz': 'Uzbek', 've': 'Venda', 'vi': 'Vietnamese', 'vo': 'Volapük', 'wa': 'Walloon', 'wo': 'Wolof', 'xh': 'Xhosa', 'yi': 'Yiddish', 'yo': 'Yoruba', 'za': 'Zhuang; Chuang', 'zu': 'Zulu'}

lang = {}


@app.route('/')
def index():
    return render_template("Main.html")


@app.route('/login')
def login():
    return render_template("Login.html")


@app.route('/signup')
def signup():
    return render_template("signup.html")


@app.route('/additional')
def additional():
    return jsonify(message="under construction")


@app.route('/translate')
def translate_two():
    return render_template("translate.html")


@app.route('/addDetails', methods=['POST'])
def addDetails():
    if request.is_json:
        lang['from'] = request.json['lang']

    return jsonify({"lang": lang['from']})


@app.route('/processUpload', methods=['POST'])
def processingUpload():
    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            flash('No file part')
            return redirect(url_for('init'))

        file = request.files['file']
        # if user does not select file, browser also
        # submit an empty part without filename

        if file.filename == '':
            flash('No selected file')
            return redirect(url_for('init'))

        if file and allowed_file(file.filename):
            # filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], file.filename))
            print(file.filename)
            # langcode = lang['from']
            # print(langcode)
            # search = lang
            # for code, og in map.items():  # for name, age in dictionary.iteritems():  (for Python 2.x)
            #     if og == search:
            #         print(code)

            processResume(lang, file.filename)

        # return render_template('displayPage.html', location=location), 200
            return flask.send_from_directory("./static/", file.filename, as_attachment=True)
        else:
            flash(f"{file.filename.split('.')[1]} is not allowed")
            return redirect(url_for('init'))


# def translate_text(text, target="en"):
#     """Translates text into the target language.

#     Target must be an ISO 639-1 language code.
#     See https://g.co/cloud/translate/v2/translate-reference#supported_languages
#     """

#     translate_client = translate.Client()

#     if isinstance(text, six.binary_type):
#         text = text.decode("utf-8")

#     # Text can also be a sequence of strings, in which case this method
#     # will return a sequence of results for each text.
#     result = translate_client.translate(
#         text, target_language=target)

#     return result["translatedText"]
#     # print(u"Text: {}".format(result["input"]))
#     # print(u"Translation: {}".format(result["translatedText"]))
#     # print(u"Detected source language: {}".format(
#     #     result["detectedSourceLanguage"]))


def translate_text(text, project_id="metal-sorter-341818"):
    """Translating Text."""

    client = translate.TranslationServiceClient()
    # print(client)

    location = "global"

    parent = f"projects/{project_id}/locations/{location}"

    # Translate text from English to French
    # Detail on supported types can be found here:
    # https://cloud.google.com/translate/docs/supported-formats
    response = client.translate_text(
        request={
            "parent": parent,
            "contents": [text],
            "mime_type": "text/plain",  # mime types: text/plain, text/html
            "source_language_code": "es",
            "target_language_code": "en",
        }
    )

    # Display the translation for each input text provided
    return response.translations[0].translated_text
#     # for translation in response.translations:
#     #     print("Translated text: {}".format(translation.translated_text))


def find_replace(paragraph_keyword, draft_keyword, paragraph, style, isBold):
    if paragraph_keyword in paragraph.text:
        # print("found")
        paragraph.text = paragraph.text.replace(
            paragraph_keyword, draft_keyword)
        paragraph.style = style
        paragraph.style.font.bold = True


def find_replace_table(paragraph_keyword, draft_keyword, cell):
    if paragraph_keyword in cell.text:
        # print("found")
        cell.text = cell.text.replace(
            paragraph_keyword, draft_keyword)


def processResume(lang, filename):
    print(lang, filename)

    filenamep = os.path.join('./upload', filename)

    langf = lang

    # upload PDF file to storage

    doc = docx.Document(filenamep)
    for i in doc.paragraphs:
        style = i.style.name
        isBold = i.style.font.bold
        # print(font)
        each = i.text
        # translate_text("en", each)
        print("Each: " + each)
        splitat = 255
        left, right = each[:splitat], each[splitat:]
        if left != "":
            print("Left: " + left)
            translation = translate_text(left)
            find_replace(left, translation, i, style, isBold)
            print(translation)

        if right != "":
            translation = translate_text(right)
            find_replace(left, translation, i, style, isBold)
            print(translation)

    location = "./static/" + filename

    doc.save(location)

    # Open a file
    # fo = open(response, "r+")
    # # str = fo.read(10)
    # # print("Read String is : ", str)
    # # Close opend file
    # fo.close()


def tryThis():

    # read document from file
    # with open('sample.txt', mode='r') as infile:
    #     for each in infile:
    #         # translate_text("en", each)
    #         print("Each: " + each)
    #         splitat = 255
    #         left, right = each[:splitat], each[splitat:]
    #         print("Left: " + left)
    #         translation = translate_text("en", left)
    #         print(translation)

    #         if right != "":
    #             translation = translate_text("en", right)
    #             print(translation)
    #         print("Right: " + right)
    # json_string = json.dumps(languages)
    # print(json_string)

    doc = docx.Document('./mexico.docx')
    for i in doc.paragraphs:
        style = i.style.name
        isBold = i.style.font.bold
        # print(font)
        each = i.text
        # translate_text("en", each)
        print("Each: " + each)
        splitat = 255
        left, right = each[:splitat], each[splitat:]
        if left != "":
            print("Left: " + left)
            translation = translate_text(left)
            find_replace(left, translation, i, style, isBold)
            print(translation)

        if right != "":
            translation = translate_text(right)
            # find_replace(left, translation, i, style, isBold)
            print(translation)
        # print("Right: " + right)
# if text it stored in a table format
# TODO seems like the ' char is not escaping properly etc

    # for a in doc.tables:
    #     for b in a._cells:
    #         each = b.text
    #         print("Each: " + each)
    #         splitat = 255
    #         left, right = each[:splitat], each[splitat:]
    #         if left != "":
    #             print("Left: " + left)
    #             translation = translate_text("it", left)
    #             find_replace(left, translation, b)
    #             print(translation)

    #         if right != "":
    #             translation = translate_text("it", right)
    #             find_replace(left, translation, b)
    #             print(translation)
    #         print("Right: " + right)
    # location = "./static/" + file.filename

    doc.save('./result.docx')


# tryThis()


def allowed_file(filename):
    print(filename.split('.', 1))
    return '.' in filename and \
           filename.split('.', 1)[1].lower() in ALLOWED_EXTENSIONS


if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True)
