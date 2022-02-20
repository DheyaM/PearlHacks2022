# import docx
# from docx import Document
# import json
# from docx.oxml.shared import OxmlElement
# from docx.oxml.ns import qn

# from google.cloud import translate

# from flask import Flask, render_template,  request, jsonify, make_response, session, flash, redirect
# from flask_pymongo import PyMongo
# import os

# app = Flask(__name__, static_url_path='')

# # import six
# # from google.cloud import translate_v2 as translate


# @app.route('/')
# def index():
#     return render_template("Main.html")


# def translate_text(text, project_id="metal-sorter-341818"):
#     """Translating Text."""

#     client = translate.TranslationServiceClient()
#     print(client)

#     location = "global"

#     parent = f"projects/{project_id}/locations/{location}"

#     # Translate text from English to French
#     # Detail on supported types can be found here:
#     # https://cloud.google.com/translate/docs/supported-formats
#     response = client.translate_text(
#         request={
#             "parent": parent,
#             "contents": [text],
#             "mime_type": "text/plain",  # mime types: text/plain, text/html
#             "source_language_code": "es",
#             "target_language_code": "en",
#         }
#     )

#     # Display the translation for each input text provided
#     return response.translations[0].translated_text
#     # for translation in response.translations:
#     #     print("Translated text: {}".format(translation.translated_text))


# def find_replace(paragraph_keyword, draft_keyword, paragraph, style, isBold):
#     if paragraph_keyword in paragraph.text:
#         # print("found")
#         paragraph.text = paragraph.text.replace(
#             paragraph_keyword, draft_keyword)
#         paragraph.style = style
#         paragraph.style.font.bold = True


# def processResume():

#     doc = docx.Document('./mexico.docx')
#     for i in doc.paragraphs:
#         style = i.style.name
#         isBold = i.style.font.bold
#         # print(font)
#         each = i.text
#         # translate_text("en", each)
#         print("Each: " + each)
#         splitat = 255
#         left, right = each[:splitat], each[splitat:]
#         if left != "":
#             print("Left: " + left)
#             translation = translate_text(left)
#             find_replace(left, translation, i, style, isBold)
#             print(translation)

#         if right != "":
#             translation = translate_text(right)
#             find_replace(left, translation, i, style, isBold)
#             print(translation)
#         # print("Right: " + right)
#     # translation = translate_text("Bueno")
#     print(translation)


# processResume()


# # TODO allow the user to overwrite the language
# if __name__ == "__main__":
#     app.run(host='0.0.0.0', debug=True)


import pymongo

# https://www.fullstackpython.com/flask-helpers-make-response-examples.html
from flask import Flask, render_template,  request, jsonify, make_response, session, flash, redirect
from flask_pymongo import PyMongo
import flask
import os

import docx
from docx import Document
import json
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from google.cloud import translate

app = Flask(__name__, static_url_path='')

UPLOAD_FOLDER = './upload'
ALLOWED_EXTENSIONS = {'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

lang = {}


@app.route('/')
def index():
    return render_template("Main.html")


@app.route('/login')
def login():
    return render_template("Login.html")


@app.route('/signup')
def signup():
    return render_template("signup.html")


@app.route('/additional')
def additional():
    return jsonify(message="under construction")


@app.route('/translate')
def translate_two():
    return render_template("translate.html")


# @app.route('/addDetails', methods=['POST'])
# def addDetails():
#     if request.is_json:
#         user['name'] = request.json['name']
#         user['email'] = request.json['email']
#         user['school'] = request.json['school']


@app.route('/processUpload', methods=['POST'])
def processingUpload():
    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            flash('No file part')
            return redirect(url_for('init'))

        file = request.files['file']
        # if user does not select file, browser also
        # submit an empty part without filename

        if file.filename == '':
            flash('No selected file')
            return redirect(url_for('init'))

        if file and allowed_file(file.filename):
            # filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], file.filename))
            print(file.filename)

            processResume(lang, file.filename)

        # return render_template('displayPage.html', location=location), 200
            return flask.send_from_directory("./static/", file.filename, as_attachment=True)
        else:
            flash(f"{file.filename.split('.')[1]} is not allowed")
            return redirect(url_for('init'))


def translate_text(text, project_id="metal-sorter-341818"):
    """Translating Text."""

    client = translate.TranslationServiceClient()
    # print(client)

    location = "global"

    parent = f"projects/{project_id}/locations/{location}"

    # Translate text from English to French
    # Detail on supported types can be found here:
    # https://cloud.google.com/translate/docs/supported-formats
    response = client.translate_text(
        request={
            "parent": parent,
            "contents": [text],
            "mime_type": "text/plain",  # mime types: text/plain, text/html
            "source_language_code": "es",
            "target_language_code": "en",
        }
    )

    # Display the translation for each input text provided
    return response.translations[0].translated_text
    # for translation in response.translations:
    #     print("Translated text: {}".format(translation.translated_text))


def find_replace(paragraph_keyword, draft_keyword, paragraph, style, isBold):
    if paragraph_keyword in paragraph.text:
        # print("found")
        paragraph.text = paragraph.text.replace(
            paragraph_keyword, draft_keyword)
        paragraph.style = style
        paragraph.style.font.bold = True


def find_replace_table(paragraph_keyword, draft_keyword, cell):
    if paragraph_keyword in cell.text:
        # print("found")
        cell.text = cell.text.replace(
            paragraph_keyword, draft_keyword)


def processResume(lang, filename):
    print(lang, filename)

    filenamep = os.path.join('./upload', filename)

    langf = lang

    # upload PDF file to storage

    doc = docx.Document(filenamep)
    for i in doc.paragraphs:
        style = i.style.name
        isBold = i.style.font.bold
        # print(font)
        each = i.text
        # translate_text("en", each)
        print("Each: " + each)
        splitat = 255
        left, right = each[:splitat], each[splitat:]
        if left != "":
            print("Left: " + left)
            translation = translate_text(left)
            find_replace(left, translation, i, style, isBold)
            print(translation)

        if right != "":
            translation = translate_text(right)
            find_replace(left, translation, i, style, isBold)
            print(translation)

    location = "./static/" + filename

    doc.save(location)

    # Open a file
    # fo = open(response, "r+")
    # # str = fo.read(10)
    # # print("Read String is : ", str)
    # # Close opend file
    # fo.close()


def tryThis():

    # read document from file
    # with open('sample.txt', mode='r') as infile:
    #     for each in infile:
    #         # translate_text("en", each)
    #         print("Each: " + each)
    #         splitat = 255
    #         left, right = each[:splitat], each[splitat:]
    #         print("Left: " + left)
    #         translation = translate_text("en", left)
    #         print(translation)

    #         if right != "":
    #             translation = translate_text("en", right)
    #             print(translation)
    #         print("Right: " + right)
    # json_string = json.dumps(languages)
    # print(json_string)

    doc = docx.Document('./mexico.docx')
    for i in doc.paragraphs:
        style = i.style.name
        isBold = i.style.font.bold
        # print(font)
        each = i.text
        # translate_text("en", each)
        print("Each: " + each)
        splitat = 255
        left, right = each[:splitat], each[splitat:]
        if left != "":
            print("Left: " + left)
            translation = translate_text(left)
            find_replace(left, translation, i, style, isBold)
            print(translation)

        if right != "":
            translation = translate_text(right)
            # find_replace(left, translation, i, style, isBold)
            print(translation)
        # print("Right: " + right)
# if text it stored in a table format
# TODO seems like the ' char is not escaping properly etc

    # for a in doc.tables:
    #     for b in a._cells:
    #         each = b.text
    #         print("Each: " + each)
    #         splitat = 255
    #         left, right = each[:splitat], each[splitat:]
    #         if left != "":
    #             print("Left: " + left)
    #             translation = translate_text("it", left)
    #             find_replace(left, translation, b)
    #             print(translation)

    #         if right != "":
    #             translation = translate_text("it", right)
    #             find_replace(left, translation, b)
    #             print(translation)
    #         print("Right: " + right)
    # location = "./static/" + file.filename

    doc.save('./result.docx')


# tryThis()


def allowed_file(filename):
    print(filename.split('.', 1))
    return '.' in filename and \
           filename.split('.', 1)[1].lower() in ALLOWED_EXTENSIONS


if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True)
